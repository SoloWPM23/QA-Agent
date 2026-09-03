"""Tests for the test suite DOCX generator."""

from __future__ import annotations

from docx import Document

from app.input.test_suite_generator import generate_test_suite_docx
from app.llm.openapi_schemas import GeneratedTestCase


def test_generate_test_suite_docx_creates_file(tmp_path):
    cases = [
        GeneratedTestCase(
            id="TC-001",
            judul="Health check",
            deskripsi="Memastikan endpoint health aktif.",
            method="GET",
            path="/health",
            expected_status_code="200",
            expected_schema="status (string)",
        )
    ]
    output = str(tmp_path / "suite.docx")
    result = generate_test_suite_docx(cases, output_path=output)

    doc = Document(result)
    assert len(doc.tables) == 1
    text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
    assert "Test Case TC-001" in text
    assert "Health check" in text
    assert "/health" in text
    assert "200" in text
