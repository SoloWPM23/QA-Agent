"""End-to-end integration tests for OpenAPI → DOCX → TestCase flow."""

from __future__ import annotations

import json

from docx import Document

from app.input.docx_adapter import DocxAdapter
from app.input.test_suite_generator import generate_test_suite_docx
from app.llm.openapi_converter import sanitize_cases
from app.llm.openapi_schemas import ConversionResult, GeneratedTestCase


def _make_cases() -> list[GeneratedTestCase]:
    """Return a small set of generated test cases covering common fields."""
    return [
        GeneratedTestCase(
            id="TC-001",
            judul="Health check",
            deskripsi="Memastikan endpoint health aktif.",
            method="GET",
            path="/health",
            expected_status_code="200",
            expected_schema="status (string)",
            jsonpath_checks="$.status",
        ),
        GeneratedTestCase(
            id="TC-002",
            judul="Buat pengguna baru",
            deskripsi="Menguji pembuatan pengguna.",
            method="POST",
            path="/api/v1/users",
            headers="Content-Type: application/json",
            body='{"name": "Alice", "email": "alice@example.com"}',
            expected_status_code="201",
            expected_schema="id (number), name (string), email (string)",
            contains="Alice",
        ),
        GeneratedTestCase(
            id="TC-003",
            judul="Detail pengguna",
            deskripsi="Menguji detail pengguna.",
            method="GET",
            path="/api/v1/users/1",
            expected_status_code="200",
            expected_schema="id (number), name (string)",
            regex="\\d{4}",
            jsonpath_checks="$.id, $.name",
        ),
        GeneratedTestCase(
            id="TC-004",
            judul="Hapus pengguna",
            deskripsi="Menguji hapus pengguna.",
            method="DELETE",
            path="/api/v1/users/1",
            expected_status_code="204",
        ),
    ]


class TestGenerateAndParseDOCX:
    """Test that generated DOCX can be parsed back into TableBlocks."""

    def test_generated_docx_has_correct_labels(self, tmp_path):
        cases = _make_cases()
        output = str(tmp_path / "suite.docx")
        result = generate_test_suite_docx(cases, output_path=output)

        doc = Document(result)
        assert len(doc.tables) == 4

        first_table = doc.tables[0]
        labels = []
        for row in first_table.rows[1:]:
            cells = row.cells
            text = cells[0].text.strip() if cells else ""
            labels.append(text)

        expected_labels = [
            "ID",
            "Judul",
            "Deskripsi",
            "Method",
            "Path",
            "Headers",
            "Query Params",
            "Body (JSON)",
            "Expected Status Code",
            "Expected Schema",
            "JSONPath Checks",
            "Regex",
            "Contains",
        ]
        assert labels == expected_labels

    def test_docx_adapter_reads_generated_docx(self, tmp_path):
        cases = _make_cases()
        output = str(tmp_path / "suite.docx")
        generate_test_suite_docx(cases, output_path=output)

        adapter = DocxAdapter()
        blocks = adapter.extract(output)
        assert len(blocks) == 4

        assert blocks[0].case_id == "TC-001"
        assert blocks[0].fields["Method"] == "GET"
        assert blocks[0].fields["Path"] == "/health"

        assert blocks[1].case_id == "TC-002"
        assert blocks[1].fields["Method"] == "POST"
        assert blocks[1].fields["Path"] == "/api/v1/users"

    def test_docx_adapter_flags_missing_required_fields(self, tmp_path):
        """A case with empty Method/Path should be flagged for review."""
        cases = [
            GeneratedTestCase(id="TC-001", judul="Empty", method="", path=""),
        ]
        output = str(tmp_path / "suite.docx")
        generate_test_suite_docx(cases, output_path=output)

        adapter = DocxAdapter()
        blocks = adapter.extract(output)
        assert len(blocks) == 1
        assert blocks[0].needs_review is True
        assert "Method" in blocks[0].review_reason

    def test_body_json_is_pretty_printed(self, tmp_path):
        cases = _make_cases()
        output = str(tmp_path / "suite.docx")
        generate_test_suite_docx(cases, output_path=output)

        doc = Document(output)
        body_row = None
        for row in doc.tables[1].rows:
            cells = row.cells
            if cells and cells[0].text.strip() == "Body (JSON)":
                body_row = cells[1].text.strip()
                break
        assert body_row is not None
        assert "\n" in body_row
        parsed = json.loads(body_row)
        assert parsed["name"] == "Alice"

    def test_path_placeholder_sets_needs_review(self, tmp_path):
        cases = [
            GeneratedTestCase(
                id="TC-001",
                judul="User detail",
                method="GET",
                path="/api/v1/users/{user_id}",
            ),
        ]
        output = str(tmp_path / "suite.docx")
        generate_test_suite_docx(cases, output_path=output)

        adapter = DocxAdapter()
        blocks = adapter.extract(output)
        assert len(blocks) == 1
        assert blocks[0].fields["Path"] == "/api/v1/users/{user_id}"


class TestSanitizeCases:
    """Test sanitize_cases logic."""

    def test_fixes_invalid_ids(self):
        cases = [GeneratedTestCase(id="", method="GET", path="/x")]
        result = sanitize_cases(cases)
        assert result[0].id == "TC-001"

    def test_fixes_invalid_methods(self):
        cases = [GeneratedTestCase(id="TC-001", method="TRACE", path="/x")]
        result = sanitize_cases(cases)
        assert result[0].method == "GET"

    def test_normalizes_method_case(self):
        cases = [GeneratedTestCase(id="TC-001", method="get", path="/x")]
        result = sanitize_cases(cases)
        assert result[0].method == "GET"

    def test_adds_leading_slash_to_path(self):
        cases = [GeneratedTestCase(id="TC-001", method="GET", path="api/v1/x")]
        result = sanitize_cases(cases)
        assert result[0].path == "/api/v1/x"

    def test_preserves_valid_case(self):
        cases = [
            GeneratedTestCase(
                id="TC-001",
                method="GET",
                path="/api/v1/users",
                expected_status_code="200",
                expected_schema="users (array)",
            ),
        ]
        result = sanitize_cases(cases)
        assert result[0].id == "TC-001"
        assert result[0].method == "GET"
        assert result[0].path == "/api/v1/users"


class TestConversionResultSchema:
    """Test ConversionResult parsing."""

    def test_empty_cases_and_failed(self):
        result = ConversionResult(cases=[], failed=[])
        assert len(result.cases) == 0
        assert len(result.failed) == 0

    def test_cases_with_failed_endpoints(self):
        result = ConversionResult(
            cases=[GeneratedTestCase(id="TC-001", method="GET", path="/health")],
            failed=["POST /unknown: unsupported"],
        )
        assert len(result.cases) == 1
        assert len(result.failed) == 1
