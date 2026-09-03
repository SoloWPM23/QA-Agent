"""Generate a test suite DOCX following the standard template."""

from __future__ import annotations

import json

from docx import Document
from docx.shared import Inches

from app.llm.openapi_schemas import GeneratedTestCase


def _pretty_json(value: str) -> str:
    """If value is valid JSON, pretty-print it; otherwise return as-is."""
    try:
        obj = json.loads(value)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except (json.JSONDecodeError, TypeError):
        return value


def generate_test_suite_docx(
    cases: list[GeneratedTestCase],
    title: str = "Test Suite",
    output_path: str = "test_suite.docx",
) -> str:
    """Generate a .docx file from generated test cases."""
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph("Dokumen ini mengikuti template standar: 1 tabel per test case.")

    for case in cases:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.5)

        header_cells = table.rows[0].cells
        header_cells[0].text = f"Test Case {case.id}"
        header_cells[1].text = f"Test Case {case.id}"

        body_value = _pretty_json(case.body)

        rows = [
            ("ID", case.id),
            ("Judul", case.judul),
            ("Deskripsi", case.deskripsi),
            ("Method", case.method),
            ("Path", case.path),
            ("Headers", case.headers),
            ("Query Params", case.query_params),
            ("Body (JSON)", body_value),
            ("Expected Status Code", case.expected_status_code),
            ("Expected Schema", case.expected_schema),
            ("JSONPath Checks", case.jsonpath_checks),
            ("Regex", case.regex),
            ("Contains", case.contains),
        ]
        for label, value in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph()

    doc.save(output_path)
    return output_path
