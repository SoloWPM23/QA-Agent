"""Generate a test-suite DOCX that matches the mock API in mock/api.py.

Run from the project root:

    venv/Scripts/python.exe mock/generate_test_suite.py

Output: ``mock/test_suite.docx`` (ready to upload via web UI or CLI).
"""

from __future__ import annotations

from docx import Document
from docx.shared import Inches

TEST_CASES = [
    {
        "id": "TC-001",
        "judul": "Health endpoint mengembalikan status ok",
        "deskripsi": "Memastikan endpoint health aktif dan merespons dengan benar.",
        "method": "GET",
        "path": "/health",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "200",
        "schema": "status (string)",
        "jsonpath": "$.status",
        "regex": "",
        "contains": "ok",
    },
    {
        "id": "TC-002",
        "judul": "List users mengembalikan daftar pengguna",
        "deskripsi": "Memastikan endpoint daftar pengguna mengembalikan array users.",
        "method": "GET",
        "path": "/api/v1/users",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "200",
        "schema": "users (array)",
        "jsonpath": "",
        "regex": "",
        "contains": "Alice",
    },
    {
        "id": "TC-003",
        "judul": "Get user detail mengembalikan data pengguna dengan id number",
        "deskripsi": "Mengambil detail pengguna ID 1.",
        "method": "GET",
        "path": "/api/v1/users/1",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "200",
        "schema": "id (number), name (string), email (string)",
        "jsonpath": "$.id",
        "regex": "",
        "contains": "alice@example.com",
    },
    {
        "id": "TC-004",
        "judul": "Create user membuat pengguna baru",
        "deskripsi": "Memastikan endpoint create user mengembalikan data pengguna yang baru dibuat.",
        "method": "POST",
        "path": "/api/v1/users",
        "headers": "Content-Type: application/json",
        "query": "(tidak ada)",
        "body": '{"name": "Charlie", "email": "charlie@example.com"}',
        "status": "201",
        "schema": "id (number), name (string), email (string)",
        "jsonpath": "$.name",
        "regex": "",
        "contains": "charlie@example.com",
    },
    {
        "id": "TC-005",
        "judul": "Update user memperbarui data pengguna",
        "deskripsi": "Memastikan endpoint update user mengembalikan data yang diperbarui.",
        "method": "PUT",
        "path": "/api/v1/users/2",
        "headers": "Content-Type: application/json",
        "query": "(tidak ada)",
        "body": '{"name": "Bobby"}',
        "status": "200",
        "schema": "id (number), name (string), email (string)",
        "jsonpath": "",
        "regex": "",
        "contains": "Bobby",
    },
    {
        "id": "TC-006",
        "judul": "Delete user menghapus pengguna",
        "deskripsi": "Memastikan endpoint delete user merespons 204 tanpa body.",
        "method": "DELETE",
        "path": "/api/v1/users/2",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "204",
        "schema": "(tidak ada)",
        "jsonpath": "",
        "regex": "",
        "contains": "",
    },
    {
        "id": "TC-007",
        "judul": "List books mengembalikan daftar buku",
        "deskripsi": "Memastikan endpoint daftar buku mengembalikan array books.",
        "method": "GET",
        "path": "/api/v1/books",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "200",
        "schema": "books (array)",
        "jsonpath": "",
        "regex": "",
        "contains": "Python 101",
    },
    {
        "id": "TC-008",
        "judul": "Get nonexistent book mengembalikan 404",
        "deskripsi": "Mengambil buku ID 99 yang tidak ada.",
        "method": "GET",
        "path": "/api/v1/books/99",
        "headers": "(tidak ada)",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "404",
        "schema": "(tidak ada)",
        "jsonpath": "",
        "regex": "",
        "contains": "",
    },
    {
        "id": "TC-009",
        "judul": "Login sukses mengembalikan access token",
        "deskripsi": "Memastikan endpoint login mengembalikan access_token dan token_type.",
        "method": "POST",
        "path": "/api/v1/login",
        "headers": "Content-Type: application/json",
        "query": "(tidak ada)",
        "body": '{"username": "admin", "password": "secret"}',
        "status": "200",
        "schema": "access_token (string), token_type (string)",
        "jsonpath": "$.access_token",
        "regex": "",
        "contains": "Bearer",
    },
    {
        "id": "TC-010",
        "judul": "Get current user dengan bearer token",
        "deskripsi": "Memastikan endpoint me mengembalikan data user saat ini.",
        "method": "GET",
        "path": "/api/v1/me",
        "headers": "Authorization: Bearer mock-token-12345",
        "query": "(tidak ada)",
        "body": "(tidak ada)",
        "status": "200",
        "schema": "id (number), name (string), email (string)",
        "jsonpath": "$.name",
        "regex": "",
        "contains": "admin@example.com",
    },
]


def _set_cell_text(cell, text: str) -> None:
    cell.text = text


def generate(path: str = "mock/test_suite.docx") -> str:
    doc = Document()
    doc.add_heading("Test Suite - Mock API", level=0)
    doc.add_paragraph("Dokumen ini mengikuti template standar: 1 tabel per test case.")

    for case in TEST_CASES:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.5)

        header_cells = table.rows[0].cells
        _set_cell_text(header_cells[0], f"Test Case {case['id']}")
        _set_cell_text(header_cells[1], f"Test Case {case['id']}")

        rows = [
            ("ID", case["id"]),
            ("Judul", case["judul"]),
            ("Deskripsi", case["deskripsi"]),
            ("Method", case["method"]),
            ("Path", case["path"]),
            ("Headers", case["headers"]),
            ("Query Params", case["query"]),
            ("Body (JSON)", case["body"]),
            ("Expected Status Code", case["status"]),
            ("Expected Schema", case["schema"]),
            ("JSONPath Checks", case["jsonpath"]),
            ("Regex", case["regex"]),
            ("Contains", case["contains"]),
        ]
        for label, value in rows:
            row_cells = table.add_row().cells
            _set_cell_text(row_cells[0], label)
            _set_cell_text(row_cells[1], value)

        doc.add_paragraph()

    doc.save(path)
    return path


if __name__ == "__main__":
    out = generate()
    print(f"Test suite tersimpan di: {out}")
